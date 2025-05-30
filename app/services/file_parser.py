import os
import logging
from io import BytesIO
from fastapi import UploadFile
from app.utils.error_handler import raise_http_exception
from app.config.settings import settings

# Initialize logger
logger = logging.getLogger(__name__)

# Set maximum file size
MAX_FILE_SIZE_MB = settings.MAX_FILE_SIZE_MB
MAX_FILE_SIZE = MAX_FILE_SIZE_MB * 1024 * 1024  # Convert MB to bytes

# Allowed file extensions
ALLOWED_EXTENSIONS = [".pdf", ".docx", ".txt", ".md"]

# Check model related libraries
try:
    import torch
    import tensorflow
    import flax
except ImportError as e:
    logger.error(f"Model libraries missing: {e.name}")
    raise_http_exception("Required model libraries are missing", code=500)


def validate_file_type(filename: str):
    # Validate file type
    if not filename:
        raise_http_exception("Filename is required", code=400)

    file_extension = os.path.splitext(filename)[1].lower()
    if file_extension not in ALLOWED_EXTENSIONS:
        supported_formats = ", ".join(ALLOWED_EXTENSIONS)
        raise_http_exception(
            f"Unsupported file type: {file_extension}. Allowed formats are: {supported_formats}.",
            code=415,
        )


def validate_file_size(file_size: int):
    # Validate file size
    if file_size > MAX_FILE_SIZE:
        raise_http_exception(f"File size exceeds {MAX_FILE_SIZE_MB} MB", code=413)


async def validate_file(file: UploadFile):
    # File validation function
    validate_file_type(file.filename)

    # FastAPI UploadFile doesn't have size attribute
    content = await file.read()
    file_size = len(content)
    await file.seek(0)

    validate_file_size(file_size)
    return file_size


async def extract_text_from_file(file: UploadFile) -> str:
    # Extract text from file
    try:
        # File validation
        await validate_file(file)

        # Read file content
        contents = await file.read()

        # Extract text based on file type
        if file.filename.lower().endswith(".pdf"):
            return extract_from_pdf(contents)
        elif file.filename.lower().endswith(".docx"):
            return extract_from_docx(contents)
        elif file.filename.lower().endswith((".txt", ".md")):
            # Decode plain text files
            return contents.decode("utf-8", errors="replace")
        else:
            raise_http_exception("Unsupported file format", code=415)

    except ValueError as e:
        # File processing error
        logger.error(f"File processing error: {str(e)}")
        raise_http_exception(f"File processing error: {str(e)}", code=400)
    except Exception as e:
        # Unexpected error
        logger.error(f"Unexpected error processing file: {str(e)}", exc_info=True)
        raise_http_exception(f"Failed to process file: {str(e)}", code=500)
    finally:
        # Reset file pointer
        try:
            await file.seek(0)
        except Exception:
            pass


def extract_from_pdf(file_bytes: bytes) -> str:
    # Extract text from PDF file
    try:
        import fitz

        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            if doc.page_count == 0:
                raise ValueError("The PDF file is empty or invalid")

            # Extract text from all pages
            text = ""
            for page in doc:
                text += page.get_text()

            return text
    except ImportError:
        logger.error("PyMuPDF (fitz) library not available")
        raise_http_exception("PDF processing is not available", code=501)
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}", exc_info=True)
        raise_http_exception(f"Failed to extract text from PDF: {str(e)}", code=500)


def extract_from_docx(file_bytes: bytes) -> str:
    # Extract text from DOCX file
    try:
        import docx

        doc = docx.Document(BytesIO(file_bytes))

        if not doc.paragraphs:
            raise ValueError("The DOCX file is empty or invalid")

        # Extract text from all paragraphs
        text = "\n".join(para.text for para in doc.paragraphs if para.text)

        return text
    except ImportError:
        logger.error("python-docx library not available")
        raise_http_exception("DOCX processing is not available", code=501)
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}", exc_info=True)
        raise_http_exception(f"Failed to extract text from DOCX: {str(e)}", code=500)
