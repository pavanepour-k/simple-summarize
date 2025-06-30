"""File upload and processing tests."""
from __future__ import annotations

import io
import tempfile
from pathlib import Path

import pytest
from fastapi import status


class TestFileUpload:
    """File upload endpoint tests."""
    
    def test_pdf_upload_success(self, client, auth_headers, sample_pdf_file):
        response = client.post(
            "/file/upload",
            files={"file": sample_pdf_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert data["filename"] == "test.pdf"
        assert data["content_type"] == "application/pdf"
        assert data["size"] > 0
        assert data["text_length"] > 0
    
    def test_docx_upload_success(self, client, auth_headers, sample_docx_file):
        response = client.post(
            "/file/upload",
            files={"file": sample_docx_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert data["filename"] == "test.docx"
        assert "wordprocessingml" in data["content_type"]
        assert data["text_length"] > 0
    
    def test_txt_upload_success(self, client, auth_headers, sample_text_file):
        response = client.post(
            "/file/upload",
            files={"file": sample_text_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert data["filename"] == "test.txt"
        assert data["content_type"] == "text/plain"
    
    def test_markdown_upload_success(self, client, auth_headers):
        md_content = b"# Markdown\n\n**Bold** text"
        md_file = ("test.md", md_content, "text/markdown")
        
        response = client.post(
            "/file/upload",
            files={"file": md_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
    
    def test_file_size_limit_validation(self, client, auth_headers, large_file):
        response = client.post(
            "/file/upload",
            files={"file": large_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_413_REQUEST_ENTITY_TOO_LARGE
        assert "exceeds maximum size" in response.json()["detail"]
    
    def test_unsupported_file_type(self, client, auth_headers):
        exe_file = ("test.exe", b"binary content", "application/octet-stream")
        
        response = client.post(
            "/file/upload",
            files={"file": exe_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "Invalid file type" in response.json()["detail"]
    
    def test_missing_file(self, client, auth_headers):
        response = client.post(
            "/file/upload",
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_422_UNPROCESSABLE_ENTITY
    
    def test_empty_file(self, client, auth_headers):
        empty_file = ("empty.txt", b"", "text/plain")
        
        response = client.post(
            "/file/upload",
            files={"file": empty_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_400_BAD_REQUEST
    
    def test_filename_sanitization(self, client, auth_headers):
        content = b"Test content"
        unsafe_file = ("../../../etc/passwd", content, "text/plain")
        
        response = client.post(
            "/file/upload",
            files={"file": unsafe_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
        assert "../" not in response.json()["filename"]
    
    def test_concurrent_file_uploads(self, client, auth_headers):
        files = [
            ("file1.txt", b"Content 1", "text/plain"),
            ("file2.txt", b"Content 2", "text/plain"),
            ("file3.txt", b"Content 3", "text/plain")
        ]
        
        responses = []
        for file_data in files:
            response = client.post(
                "/file/upload",
                files={"file": file_data},
                headers=auth_headers
            )
            responses.append(response)
        
        for response in responses:
            assert response.status_code == status.HTTP_200_OK
    
    def test_special_characters_in_filename(self, client, auth_headers):
        content = b"Test content"
        special_file = ("file@#$%.txt", content, "text/plain")
        
        response = client.post(
            "/file/upload",
            files={"file": special_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
    
    def test_unicode_filename(self, client, auth_headers):
        content = b"Test content"
        unicode_file = ("文件名.txt", content, "text/plain")
        
        response = client.post(
            "/file/upload",
            files={"file": unicode_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK


class TestFileSummarization:
    """File-based summarization tests."""
    
    def test_pdf_summarization(self, client, auth_headers, sample_pdf_file):
        response = client.post(
            "/file/summarize?style=general&length=medium",
            files={"file": sample_pdf_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert "summary" in data
        assert data["style"] == "general"
    
    def test_docx_summarization(self, client, auth_headers, sample_docx_file):
        response = client.post(
            "/file/summarize?style=problem_solver&length=short",
            files={"file": sample_docx_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert data["style"] == "problem_solver"
    
    def test_text_extraction_accuracy(self, client, auth_headers):
        known_content = "This is a specific test content for accuracy validation."
        txt_file = ("accuracy.txt", known_content.encode(), "text/plain")
        
        response = client.post(
            "/file/summarize",
            files={"file": txt_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert data["original_length"] == len(known_content)
    
    def test_language_detection_from_file(self, client, auth_headers):
        spanish_content = "Este es un texto en español para probar la detección."
        txt_file = ("spanish.txt", spanish_content.encode(), "text/plain")
        
        response = client.post(
            "/file/summarize",
            files={"file": txt_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
    
    def test_corrupted_pdf_handling(self, client, auth_headers):
        corrupted_pdf = ("corrupted.pdf", b"%PDF-1.4\nCorrupted", "application/pdf")
        
        response = client.post(
            "/file/summarize",
            files={"file": corrupted_pdf},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_400_BAD_REQUEST
    
    def test_password_protected_pdf(self, client, auth_headers):
        import fitz
        
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Protected content")
        
        temp_file = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        doc.save(temp_file.name, encryption=fitz.PDF_ENCRYPT_AES_256, user_pw="password")
        doc.close()
        
        with open(temp_file.name, "rb") as f:
            protected_pdf = ("protected.pdf", f.read(), "application/pdf")
        
        response = client.post(
            "/file/summarize",
            files={"file": protected_pdf},
            headers=auth_headers
        )
        assert response.status_code in [status.HTTP_400_BAD_REQUEST, status.HTTP_500_INTERNAL_SERVER_ERROR]
        
        Path(temp_file.name).unlink()
    
    def test_summarization_consistency(self, client, auth_headers):
        content = "Test content for consistency check " * 50
        txt_file = ("test.txt", content.encode(), "text/plain")
        
        responses = []
        for _ in range(3):
            response = client.post(
                "/file/summarize?style=general&length=short",
                files={"file": txt_file},
                headers=auth_headers
            )
            responses.append(response.json())
        
        summaries = [r["summary"] for r in responses]
        assert len(set(summaries)) == 1
    
    def test_file_summarization_with_all_params(self, client, auth_headers):
        txt_file = ("test.txt", b"Test content", "text/plain")
        
        response = client.post(
            "/file/summarize?style=emotion_focused&length=long",
            files={"file": txt_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert data["style"] == "emotion_focused"


class TestFileProcessingEdgeCases:
    """Edge cases for file processing."""
    
    def test_pdf_with_images_only(self, client, auth_headers):
        import fitz
        import io
        from PIL import Image
        
        doc = fitz.open()
        page = doc.new_page()
        
        img = Image.new('RGB', (100, 100), color='red')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        
        page.insert_image(fitz.Rect(0, 0, 100, 100), stream=img_bytes.read())
        
        pdf_bytes = doc.tobytes()
        doc.close()
        
        image_pdf = ("image_only.pdf", pdf_bytes, "application/pdf")
        
        response = client.post(
            "/file/summarize",
            files={"file": image_pdf},
            headers=auth_headers
        )
        assert response.status_code in [status.HTTP_200_OK, status.HTTP_400_BAD_REQUEST]
    
    def test_docx_with_tables(self, client, auth_headers):
        from docx import Document
        
        doc = Document()
        doc.add_heading('Document with Tables', 0)
        
        table = doc.add_table(rows=3, cols=3)
        for i in range(3):
            for j in range(3):
                table.cell(i, j).text = f"Cell {i},{j}"
        
        doc.add_paragraph("Text after table")
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        table_docx = ("table.docx", docx_bytes.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        response = client.post(
            "/file/summarize",
            files={"file": table_docx},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
    
    def test_mixed_encoding_text_file(self, client, auth_headers):
        mixed_content = "ASCII text plus UTF-8: café, naïve"
        txt_file = ("mixed.txt", mixed_content.encode('utf-8'), "text/plain; charset=utf-8")
        
        response = client.post(
            "/file/summarize",
            files={"file": txt_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_200_OK
    
    def test_binary_file_rejection(self, client, auth_headers):
        binary_data = bytes(range(256))
        bin_file = ("data.bin", binary_data, "application/octet-stream")
        
        response = client.post(
            "/file/summarize",
            files={"file": bin_file},
            headers=auth_headers
        )
        assert response.status_code == status.HTTP_400_BAD_REQUEST